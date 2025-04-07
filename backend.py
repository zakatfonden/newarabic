# backend.py (DocAI Extraction, Gemini Processing, RTL Word Doc Creation & Merging)

import io
import google.generativeai as genai # For Gemini API calls
from docx import Document # For creating/reading .docx files
from docxcompose.composer import Composer # For merging .docx files
from docx.shared import Pt, Inches # For setting font size, potentially margins
from docx.enum.text import WD_ALIGN_PARAGRAPH # For setting paragraph alignment
# --- ADD BACK OXML IMPORTS for reliable RTL font setting ---
from docx.oxml.ns import qn # Qualified name helper for XML namespaces
from docx.oxml import OxmlElement # For creating/manipulating OpenXML elements
# ---
import logging # For logging backend operations
import os # For environment variables (credentials)
import json # For parsing credentials from secrets

# --- Import Google Cloud Document AI Client Library ---
from google.cloud import documentai
# --- Import Streamlit for accessing secrets ---
import streamlit as st
# ---

# Configure basic logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(module)s - %(message)s')

# --- Runtime Credentials Setup for Document AI ---
# Attempts to configure Google Cloud credentials, needed for Document AI API calls.
# Priority:
# 1. Streamlit Secret 'GOOGLE_CREDENTIALS_JSON' (written to temp file)
# 2. Environment variable 'GOOGLE_APPLICATION_CREDENTIALS' (path to existing file)
CREDENTIALS_FILENAME = "google_credentials.json" # Temp filename if using secrets
_credentials_configured = False # Global flag to track success

# 1. Try Streamlit secrets
if "GOOGLE_CREDENTIALS_JSON" in st.secrets:
    logging.info("Found GOOGLE_CREDENTIALS_JSON in Streamlit Secrets. Setting up credentials file.")
    try:
        # Read the JSON string from secrets
        credentials_json_content_from_secrets = st.secrets["GOOGLE_CREDENTIALS_JSON"]
        # logging.info(f"Read {len(credentials_json_content_from_secrets)} characters from secret.") # Verbose

        # Check if the secret is empty
        if not credentials_json_content_from_secrets.strip():
            logging.error("GOOGLE_CREDENTIALS_JSON secret is empty.")
            _credentials_configured = False
        else:
            file_written_successfully = False
            try:
                # Attempt to parse and clean potential newline issues in the private key
                # JSON strings in secrets might have escaped newlines (\n) instead of actual newlines.
                cleaned_content = credentials_json_content_from_secrets
                try:
                    temp_data = json.loads(credentials_json_content_from_secrets)
                    # Check if 'private_key' exists and is a string
                    if 'private_key' in temp_data and isinstance(temp_data['private_key'], str):
                        original_pk = temp_data['private_key']
                        # Replace common escaped newline representations with actual newlines
                        cleaned_pk = original_pk.replace('\\r\\n', '\n').replace('\\r', '\n').replace('\\n', '\n')
                        if cleaned_pk != original_pk:
                            logging.warning("Attempted to clean newline characters from private_key string in credentials.")
                            temp_data['private_key'] = cleaned_pk # Update the dictionary
                        # Re-serialize the potentially modified dictionary for consistent formatting
                        cleaned_content = json.dumps(temp_data, indent=2)
                    else:
                         # If no private key or not a string, just format the parsed data
                         cleaned_content = json.dumps(temp_data, indent=2)
                except json.JSONDecodeError:
                    # If initial parsing fails, log a warning and use the raw string
                    logging.warning("Initial JSON parse for cleaning credentials failed. Writing raw secret string (may cause auth issues).")
                    cleaned_content = credentials_json_content_from_secrets # Fallback to raw content

                # Write the (potentially cleaned) credentials content to the temporary file
                with open(CREDENTIALS_FILENAME, "w", encoding='utf-8') as f:
                    f.write(cleaned_content)
                # logging.info(f"Successfully wrote credentials to {CREDENTIALS_FILENAME}") # Verbose
                file_written_successfully = True
            except Exception as write_err:
                # Catch errors during file writing
                logging.error(f"CRITICAL Error writing credentials file '{CREDENTIALS_FILENAME}': {write_err}", exc_info=True)
                _credentials_configured = False

            # If the file was written successfully, set the environment variable and validate
            if file_written_successfully:
                try:
                    # Check if the file actually exists after writing
                    if os.path.exists(CREDENTIALS_FILENAME):
                        # Set the environment variable that Google Cloud libraries look for
                        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_FILENAME
                        logging.info(f"GOOGLE_APPLICATION_CREDENTIALS set to: {CREDENTIALS_FILENAME}")

                        # Optional: Attempt a quick client initialization to validate credentials early
                        try:
                            # Use a common region endpoint for the test
                            test_opts = {"api_endpoint": "us-documentai.googleapis.com"}
                            test_client = documentai.DocumentProcessorServiceClient(client_options=test_opts)
                            logging.info("Credential file seems valid (Document AI client initialized successfully).")
                            _credentials_configured = True # Mark as configured ONLY if validation succeeds
                        except Exception as client_init_err:
                            # Catch errors during client initialization (often indicates bad credentials)
                            logging.error(f"Credential validation failed: Could not initialize Document AI client. Error: {client_init_err}", exc_info=True)
                            _credentials_configured = False
                    else:
                        # Should not happen if writing succeeded, but check anyway
                        logging.error(f"Credentials file {CREDENTIALS_FILENAME} not found after attempting to write it.")
                        _credentials_configured = False
                except Exception as env_err:
                    # Catch errors related to setting the environment variable
                    logging.error(f"Error setting or checking credentials environment variable: {env_err}", exc_info=True)
                    _credentials_configured = False

    except Exception as e:
        # Catch errors reading the secret itself
        logging.error(f"CRITICAL Error reading/processing GOOGLE_CREDENTIALS_JSON secret: {e}", exc_info=True)
        _credentials_configured = False

# 2. Fallback: Check if the environment variable is set externally (e.g., in local dev, Docker)
elif "GOOGLE_APPLICATION_CREDENTIALS" in os.environ:
    logging.info("Using GOOGLE_APPLICATION_CREDENTIALS environment variable set externally.")
    ext_cred_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
    # Check if the path specified in the env var exists
    if ext_cred_path and os.path.exists(ext_cred_path):
        logging.info(f"External credentials file found at: {ext_cred_path}")
        _credentials_configured = True # Assume valid if path exists
    else:
        logging.error(f"External GOOGLE_APPLICATION_CREDENTIALS path not found or invalid: '{ext_cred_path}'")
        _credentials_configured = False

# 3. If neither secret nor env var is found
else:
    logging.warning("Google Cloud Credentials NOT found: Set 'GOOGLE_CREDENTIALS_JSON' Streamlit secret or 'GOOGLE_APPLICATION_CREDENTIALS' environment variable.")
    _credentials_configured = False
# --- END: Runtime Credentials Setup ---


# --- Function to extract text using Document AI ---
def extract_text_with_docai(file_obj, project_id: str, location: str, processor_id: str):
    """
    Extracts text from a PDF or image file object using Google Cloud Document AI.

    Args:
        file_obj: A file-like object (e.g., from st.file_uploader) containing the document.
        project_id: Google Cloud Project ID.
        location: Region of the Document AI processor (e.g., 'us', 'eu').
        processor_id: The ID of the Document AI processor to use.

    Returns:
        A string containing the extracted text on success.
        An error message string starting with 'Error:' on failure.
    """
    global _credentials_configured # Check if credentials setup was successful
    if not _credentials_configured:
        logging.error("Document AI call skipped: Credentials not configured.")
        return "Error: Document AI authentication failed (Credentials not configured)."
    # Check if required configuration parameters are provided
    if not all([project_id, location, processor_id]):
        logging.error("Document AI call skipped: Missing configuration.")
        return "Error: Missing Document AI configuration (Project ID, Location, or Processor ID)."

    try:
        # Configure the Document AI client to use the correct regional endpoint
        opts = {"api_endpoint": f"{location}-documentai.googleapis.com"}
        client = documentai.DocumentProcessorServiceClient(client_options=opts)

        # Construct the full processor resource name path
        name = client.processor_path(project_id, location, processor_id)

        # Read the file content from the file object
        file_obj.seek(0) # Ensure reading from the start of the stream
        image_content = file_obj.read()

        # --- Determine mime_type robustly ---
        # Prefer the 'type' attribute set by Streamlit uploader if available
        mime_type = getattr(file_obj, 'type', None)
        if not mime_type:
            # Fallback: Guess MIME type from the filename extension
            ext = os.path.splitext(file_obj.name)[1].lower()
            mime_map = {
                ".pdf": "application/pdf", ".png": "image/png", ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg", ".tif": "image/tiff", ".tiff": "image/tiff",
                ".gif": "image/gif", ".bmp": "image/bmp", ".webp": "image/webp"
            }
            mime_type = mime_map.get(ext)
            # Return error if type cannot be determined
            if not mime_type:
                logging.error(f"Cannot determine MIME type for file '{file_obj.name}' from extension '{ext}'.")
                return f"Error: Cannot determine MIME type for file '{file_obj.name}'."
            logging.warning(f"Guessed MIME type '{mime_type}' from extension for {file_obj.name}")
        # --- End mime_type determination ---

        # Prepare the document content for the API request
        raw_document = documentai.RawDocument(content=image_content, mime_type=mime_type)

        # Create the process request object
        # Optional: Add skip_human_review=True if using a processor version that supports it
        request = documentai.ProcessRequest(name=name, raw_document=raw_document)

        logging.info(f"Sending request to Document AI processor: {name} for file: {file_obj.name} ({mime_type})")
        # Make the API call with a timeout (e.g., 600 seconds = 10 minutes for large docs)
        result = client.process_document(request=request, timeout=600)
        document = result.document # Get the processed document object

        # Log success and return the extracted text
        logging.info(f"Received response from Document AI for '{file_obj.name}'. Extracted text length: {len(document.text)}")
        return document.text

    except Exception as e:
        # Log and format errors for returning to the frontend
        logging.error(f"Error calling Document AI for file '{file_obj.name}': {e}", exc_info=True)
        # Try to extract more specific error details if available
        error_detail = str(e)
        if hasattr(e, 'details') and callable(e.details): # gRPC errors often have details()
            error_detail = e.details()
        elif hasattr(e, 'message'): # Other exception types might have 'message'
            error_detail = e.message
        return f"Error: Failed to process document with Document AI. Details: {error_detail}"


# --- Gemini Processing Function ---
def process_text_with_gemini(api_key: str, raw_text: str, rules_prompt: str, model_name: str):
    """
    Sends text (e.g., extracted via Document AI) to the Gemini API for processing
    based on provided rules/prompt.

    Args:
        api_key: The Google Gemini API key.
        raw_text: The input text to be processed.
        rules_prompt: Instructions for the Gemini model.
        model_name: The specific Gemini model ID to use (e.g., "gemini-1.5-flash-latest").

    Returns:
        A string containing the processed text from Gemini on success.
        An empty string "" if the input raw_text was empty or whitespace only.
        An error message string starting with 'Error:' on failure.
    """
    # --- Input Validation ---
    if not api_key:
        logging.error("Gemini call skipped: API key missing.")
        return "Error: Gemini API key is missing."
    # If input text is empty or whitespace only, skip the API call and return empty string
    if not raw_text or not raw_text.strip():
        logging.warning("Skipping Gemini call: Input text is empty.")
        return ""
    if not model_name:
        logging.error("Gemini call skipped: Model name not specified.")
        return "Error: Gemini model name not specified."
    # --- End Input Validation ---

    try:
        # Configure the Gemini client library
        genai.configure(api_key=api_key)
        logging.info(f"Initializing Gemini model: {model_name}")

        # --- Configure Safety Settings (Optional but Recommended) ---
        # Block content with medium or higher probability for these categories
        safety_settings = {
             'HARM_CATEGORY_HARASSMENT': 'BLOCK_MEDIUM_AND_ABOVE',
             'HARM_CATEGORY_HATE_SPEECH': 'BLOCK_MEDIUM_AND_ABOVE',
             'HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_MEDIUM_AND_ABOVE',
             'HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_MEDIUM_AND_ABOVE',
        }

        # --- Configure Generation Parameters (Optional) ---
        generation_config = genai.types.GenerationConfig(
             max_output_tokens=8192, # Set max output tokens (check model limits)
             temperature=0.7, # Adjust creativity (0=deterministic, >1=more creative)
             # top_p=..., top_k=... # Other sampling parameters if needed
        )

        # Initialize the generative model
        model = genai.GenerativeModel(
            model_name,
            safety_settings=safety_settings,
            generation_config=generation_config
        )

        # --- Construct the Prompt ---
        # Clearly separate instructions from the text to be processed
        full_prompt = f"{rules_prompt}\n\n---\n\n{raw_text}\n\n---\n\nProcess the text above this line:"

        logging.info(f"Sending request to Gemini model: {model_name}. Input text length: {len(raw_text)}")

        # --- Make the API Call ---
        response = model.generate_content(full_prompt)

        # --- Process the Response ---
        # Check for valid candidates in the response
        if response.candidates:
            candidate = response.candidates[0] # Usually only one candidate
            # Check if the candidate has content parts
            if candidate.content and candidate.content.parts:
                # Concatenate text from all parts (usually just one part for text models)
                processed_text = "".join(part.text for part in candidate.content.parts)
                logging.info(f"Successfully received response from Gemini ({model_name}). Output length: {len(processed_text)}")
                return processed_text # Return the successful result
            # Check if generation stopped for reasons other than normal completion
            elif candidate.finish_reason.name != "STOP":
                finish_reason = candidate.finish_reason.name
                safety_ratings = getattr(candidate, 'safety_ratings', "N/A") # Get safety ratings if available
                logging.error(f"Gemini generation ({model_name}) finished unexpectedly. Reason: {finish_reason}. Safety Ratings: {safety_ratings}")
                # Handle specific finish reasons, e.g., safety blocking
                if finish_reason == "SAFETY":
                    # Extract categories that caused the block
                    blocked_categories = [rating.category.name for rating in safety_ratings if rating.probability.name not in ['NEGLIGIBLE', 'LOW']]
                    return f"Error: Content blocked by Gemini safety filters. Categories: {', '.join(blocked_categories) or 'Unknown'}."
                else:
                    # Handle other reasons like MAX_TOKENS, RECITATION, etc.
                    return f"Error: Gemini generation failed. Reason: {finish_reason}."
            else:
                # Finished normally (STOP) but no content parts? (Should be rare)
                logging.warning(f"Gemini ({model_name}) finished normally but returned no content parts.")
                return "" # Return empty string
        # Check if the prompt itself was blocked (even if no candidates generated)
        elif response.prompt_feedback and response.prompt_feedback.block_reason.name != "BLOCK_REASON_UNSPECIFIED":
            block_reason = response.prompt_feedback.block_reason.name
            safety_ratings = response.prompt_feedback.safety_ratings
            logging.error(f"Gemini prompt blocked ({model_name}). Reason: {block_reason}. Safety Ratings: {safety_ratings}")
            blocked_categories = [rating.category.name for rating in safety_ratings if rating.probability.name not in ['NEGLIGIBLE', 'LOW']]
            return f"Error: Gemini prompt blocked by safety filters. Reason: {block_reason}. Categories: {', '.join(blocked_categories) or 'Unknown'}."
        else:
            # General case for empty or unexpected response structure
            logging.error(f"Gemini ({model_name}) returned an empty or unexpected response structure: {response}")
            return "Error: Gemini returned no valid response."

    except Exception as e:
        # Catch general exceptions during API interaction
        logging.error(f"Error interacting with Gemini API ({model_name}): {e}", exc_info=True)
        return f"Error: Failed to process text with Gemini ({model_name}). Details: {str(e)}"


# --- Create Word Document Function (with RTL Formatting) ---
def create_word_doc_from_processed_text(processed_text: str, filename: str, extraction_error: bool, gemini_error: bool):
    """
    Creates a Word document (.docx) in memory containing the provided text,
    formatted for Right-to-Left (RTL) script (e.g., Arabic, Urdu) using Arial font.
    Includes placeholder text if processing failed or text is empty.

    Args:
        processed_text: The string containing the text processed by Gemini (or empty/error).
        filename: The original filename, used for placeholder text.
        extraction_error: Boolean indicating if Document AI extraction failed.
        gemini_error: Boolean indicating if Gemini processing failed.

    Returns:
        A BytesIO stream containing the Word document, or None if a critical error occurs.
    """
    try:
        # Create a new Word document object
        document = Document()

        # --- Set Default Style for Arabic/RTL ---
        style = document.styles['Normal'] # Get the default 'Normal' style
        font = style.font
        font.name = 'Arial' # Use Arial for broad Arabic/Urdu support
        font.size = Pt(12) # Set font size (adjust as needed)
        font.rtl = True    # Enable Right-to-Left rendering for the font in the style

        # --- Apply Complex Script Font via OXML ---
        # This is crucial for ensuring Word correctly uses the font for RTL scripts.
        style_element = style.element # Get the underlying OxmlElement for the style
        # Find or create the run properties element (<w:rPr>) within the style definition
        rpr = style_element.xpath('.//w:rPr')
        if not rpr: # If <w:rPr> doesn't exist, create it
            rpr = OxmlElement('w:rPr')
            style_element.append(rpr)
        else: # If it exists, get the first one
            rpr = rpr[0]
        # Find or create the font definition element (<w:rFonts>) within run properties
        font_name_el = rpr.find(qn('w:rFonts')) # qn resolves the namespace prefix 'w:'
        if font_name_el is None: # If <w:rFonts> doesn't exist, create it
            font_name_el = OxmlElement('w:rFonts')
            rpr.append(font_name_el)
        # Set the complex script font attribute (w:cs) to 'Arial'
        # 'cs' tells Word to use this font for complex scripts like Arabic/Hebrew
        font_name_el.set(qn('w:cs'), 'Arial')
        # --- End OXML Font Setting ---

        # Set paragraph format defaults for the 'Normal' style
        p_fmt = style.paragraph_format
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Right-align paragraphs for RTL
        p_fmt.right_to_left = True               # Set paragraph direction to RTL
        # --- End Style Setup ---

        # --- Add Content or Placeholder ---
        # Check if there is valid processed text and no critical errors occurred
        if processed_text and processed_text.strip() and not extraction_error and not gemini_error:
            # Split the processed text into paragraphs (simple split by newline, adjust if needed)
            paragraphs_text = processed_text.strip().split('\n')
            for para_text in paragraphs_text:
                trimmed_para = para_text.strip()
                if trimmed_para: # Avoid adding empty paragraphs
                    # Add a new paragraph; it should inherit the 'Normal' style settings (RTL)
                    p = document.add_paragraph(trimmed_para)
                    # Optionally, explicitly set alignment/direction again (belt-and-suspenders)
                    # p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # p.paragraph_format.right_to_left = True
        else:
            # If no valid text, add placeholder text indicating the issue
            p = document.add_paragraph()
            placeholder_msg = f"[Processing issue for '{filename}': "
            if extraction_error:
                placeholder_msg += "Failed to extract text from source. "
            elif gemini_error:
                 placeholder_msg += "Failed to process extracted text with Gemini. "
            elif not processed_text or not processed_text.strip():
                 # This covers cases where extraction was ok, Gemini ran ok, but returned empty
                 placeholder_msg += "No text content after processing. "
            else: # Should ideally not be reached with current logic
                 placeholder_msg += "Unknown issue. "
            placeholder_msg += "]"
            # Add the placeholder message as a run and make it italic
            run = p.add_run(placeholder_msg)
            run.italic = True
            # Ensure the placeholder paragraph itself is also formatted RTL
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True
            run.font.rtl = True # Also mark the run as RTL
        # --- End Content/Placeholder Addition ---

        # Save the document content to a BytesIO stream (in-memory file)
        doc_stream = io.BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0) # Reset stream position to the beginning for reading
        logging.info(f"Successfully created intermediate RTL Word doc stream for '{filename}'.")
        return doc_stream # Return the stream

    except Exception as e:
        # Log errors during document creation
        logging.error(f"Error creating intermediate RTL Word document for '{filename}': {e}", exc_info=True)
        return None # Return None to indicate failure


# --- Merge Word Documents Function ---
# This function remains unchanged as it merges DOCX streams regardless of content formatting.
def merge_word_documents(doc_streams_data: list[tuple[str, io.BytesIO]]):
    """
    Merges multiple Word documents (provided as BytesIO streams) into one single document.
    Adds a page break before appending each subsequent document (after the first).

    Args:
        doc_streams_data: A list of tuples, where each tuple contains:
                          (original_filename: str, doc_stream: io.BytesIO).

    Returns:
        A BytesIO stream containing the merged Word document, or None if an error occurs
        or if the input list is empty.
    """
    # Check if there are any documents to merge
    if not doc_streams_data:
        logging.warning("No document streams provided for merging.")
        return None

    try:
        # Initialize the process with the first document in the list
        first_filename, first_stream = doc_streams_data[0]
        first_stream.seek(0) # Ensure stream is at the beginning

        # Load the first document using python-docx
        master_doc = Document(first_stream)
        # Initialize the Composer object with this first document as the base
        composer = Composer(master_doc)
        logging.info(f"Initialized merger with base document from '{first_filename}'.")

        # Append the rest of the documents (if any)
        if len(doc_streams_data) > 1:
            # Loop through the remaining documents starting from the second one (index 1)
            for i in range(1, len(doc_streams_data)):
                filename, stream = doc_streams_data[i]
                stream.seek(0) # Ensure stream is at the beginning
                logging.info(f"Preparing to append content from '{filename}' (index {i})...")

                # --- Add Page Break Before Appending ---
                # Access the underlying python-docx Document object via composer.doc
                try:
                    composer.doc.add_page_break()
                    logging.info(f"Added page break before appending '{filename}'.")
                except Exception as pb_exc:
                    # Log if adding page break fails, but attempt to continue merging
                    logging.error(f"Could not add page break before '{filename}': {pb_exc}", exc_info=True)
                # --- End Page Break Addition ---

                # Load the sub-document to be appended
                try:
                    sub_doc = Document(stream)
                    # Append the content of the sub-document to the composer's document
                    composer.append(sub_doc)
                    logging.info(f"Successfully appended content from '{filename}'.")
                except Exception as append_exc:
                     # Log if appending a specific document fails, but continue
                     logging.error(f"Failed to append document '{filename}': {append_exc}", exc_info=True)
                     pass # Continue with the next file

        # Save the final composed document to a new memory stream
        merged_stream = io.BytesIO()
        composer.save(merged_stream)
        merged_stream.seek(0) # Reset stream position for reading
        logging.info(f"Successfully merged content from {len(doc_streams_data)} source documents into final stream.")
        return merged_stream # Return the final merged document stream

    except Exception as e:
        # Catch errors during initialization, saving, or unexpected issues in the loop
        logging.error(f"Error during the overall merge process: {e}", exc_info=True)
        return None # Return None to indicate failure

# --- (End of backend.py) ---
